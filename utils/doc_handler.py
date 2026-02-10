"""
Document handler for reading and writing DOCX files
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from copy import deepcopy
import os
import shutil

def read_docx(file_path):
    """
    Read content from a DOCX file and return as plain text
    """
    try:
        doc = Document(file_path)
        full_text = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        return '\n'.join(full_text)
    except Exception as e:
        raise Exception(f"Error reading DOCX file: {str(e)}")

def write_docx(content, output_path, original_doc_path=None):
    """
    Write content to a DOCX file, preserving original formatting if available
    Args:
        content: Text content to write
        output_path: Path where the file will be saved
        original_doc_path: Path to original document to copy formatting from
    """
    try:
        if original_doc_path and os.path.exists(original_doc_path):
            # Simply copy the original document as a starting point
            shutil.copy2(original_doc_path, output_path)
            
            # Load the copied document
            doc = Document(output_path)
            
            # Split new content into paragraphs (non-empty lines)
            new_content_lines = [line for line in content.split('\n') if line.strip()]
            
            # Get original paragraphs (non-empty)
            original_paras = [p for p in doc.paragraphs if p.text.strip()]
            
            # Update text content while preserving ALL formatting
            for i, para in enumerate(original_paras):
                if i < len(new_content_lines):
                    new_text = new_content_lines[i]
                    old_text = para.text
                    
                    # Only update if text actually changed
                    if old_text != new_text:
                        # Try to preserve formatting by replacing text smartly
                        # If the paragraph has multiple runs (mixed formatting), we need to be careful
                        if len(para.runs) > 1:
                            # Multiple runs means mixed formatting (bold, colors, etc.)
                            # Try to preserve as much as possible
                            # Find the longest run and replace its text
                            longest_run = max(para.runs, key=lambda r: len(r.text))
                            
                            # Clear all run texts
                            for run in para.runs:
                                run.text = ''
                            
                            # Put new text in the run that preserves most formatting
                            longest_run.text = new_text
                        else:
                            # Single run - just update text (preserves all formatting)
                            para.runs[0].text = new_text if para.runs else para.text
                            
            # Save the modified document
            doc.save(output_path)
            
        else:
            # Create new document with basic formatting
            doc = Document()
            
            # Split content into paragraphs
            paragraphs = content.split('\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph(para_text)
                    
                    # Apply basic formatting
                    for run in para.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
        
            doc.save(output_path)
        
        return True
    except Exception as e:
        raise Exception(f"Error writing DOCX file: {str(e)}")

def create_comparison_dict(original_text, modified_text):
    """
    Create a dictionary containing original and modified text for comparison
    """
    return {
        'original': original_text,
        'modified': modified_text
    }

def get_file_extension(filename):
    """
    Get file extension from filename
    """
    return os.path.splitext(filename)[1].lower()

def validate_docx(filename):
    """
    Validate if the file is a DOCX file
    """
    valid_extensions = ['.docx', '.doc']
    return get_file_extension(filename) in valid_extensions
